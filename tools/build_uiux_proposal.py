from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/seniorlifepr/.codex/skills/suiyue-document-governance/assets/suiyue-word-template.docx")
SHOTS = Path("/tmp/edoc-ui-audit-shots")
OUTPUT = ROOT / "docs" / "歲悅_總務_待配發_eDoc_UIUX改善方案_v0.1_20260824.docx"

ORANGE = "E67817"
BROWN = "6B4A36"
INK = "4F4F4F"
MUTED = "797979"
CREAM = "FFF5E8"
PEACH = "FCE3CD"
SAGE = "97AA90"
BLUE = "E4F0F4"
ALERT = "F7E4DB"
BORDER = "D9C8B8"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_run_font(run, size: float = 12, bold: bool | None = None, color: str = INK) -> None:
    run.font.name = "Kaiti TC"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Kaiti TC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Kaiti TC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Kaiti TC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Kaiti TC")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Mm(width * 25.4 / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(cell, text: str, *, bold: bool = False, color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 12) -> None:
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=BROWN)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text))


def add_heading(doc: Document, text: str, level: int = 1, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size={1: 18, 2: 16, 3: 14}.get(level, 13), bold=True, color=ORANGE if level == 1 else BROWN if level == 2 else INK)


def add_callout(doc: Document, label: str, body: str, fill: str = CREAM) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [10772])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(f"{label}　"), bold=True, color=ORANGE)
    set_run_font(p.add_run(body))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_alt_text(inline_shape, alt: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", alt)
    doc_pr.set("descr", alt)


def add_picture_to_cell(cell, image_path: Path, width_inches: float, alt: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width_inches))
    set_alt_text(shape, alt)


def add_comparison(doc: Document, title: str, before: Path, after: Path, *, mobile: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=13, bold=True, color=BROWN)
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [5386, 5386])
    for idx, label in enumerate(("Before｜改善前", "After｜改善後")):
        cell = table.cell(0, idx)
        set_cell_fill(cell, ORANGE if idx else BROWN)
        add_text(cell, label, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    width = 1.7 if mobile else 3.45
    add_picture_to_cell(table.cell(1, 0), before, width, f"{title}改善前畫面")
    add_picture_to_cell(table.cell(1, 1), after, width, f"{title}改善後畫面")


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("文件名稱", "eDoc 電子公文用印系統 UI／UX 改善方案", "文件編號", "［待公司核定］"),
        ("類別", "A｜治理／決策資料", "版次", "v0.1"),
        ("管理單位", "總務／系統專案［待公司核定］", "狀態", "審查中"),
        ("日期", "2026 年 8 月 24 日", "機密級別", "［待公司核定］"),
        ("正式位置", "［待公司核定］", "保存年限", "［待公司核定］"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1500, 3886, 1500, 3886])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if c_idx % 2 == 0:
                set_cell_fill(cell, PEACH)
                add_text(cell, value, bold=True, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_text(cell, value)


def add_status_table(doc: Document) -> None:
    data = [
        ("P0", "首頁先工作、非先看上線技術面板", "已完成", "日常首頁僅顯示今日待辦與角色工作台"),
        ("P0", "手機內容可完整往下捲", "已完成", "17 個可進入頁面皆可垂直捲動"),
        ("P0", "手機不得整頁橫向跑版", "已完成", "17 個頁面橫向溢出均為 0"),
        ("P0", "手機主要操作區至少 44 px", "已完成", "17 個頁面僅保留原生勾選框例外"),
        ("P0", "合約用印與電子用印避免雙入口", "已完成", "主選單只保留電子用印；舊網址仍相容"),
        ("P0", "一般介面用黑體；正式公文內容用標楷體", "已完成", "表單與導覽為 Noto Sans TC，公文預覽／輸出為標楷體鏈"),
        ("驗收", "Portal→eDoc 實際使用者 5 秒內進入", "待人工實測", "需由真實公司帳號重新點入並記錄秒數"),
        ("驗收", "正式 Ready 檢查", "待外部條件", "防毒／維運 readiness 尚須正式環境確認"),
        ("驗收", "公司大小章正式圖檔與實際用印", "依需求排除", "本報告只確認介面與流程，不宣稱印章素材已完成"),
    ]
    table = doc.add_table(rows=1 + len(data), cols=4)
    set_table_geometry(table, [900, 3450, 1500, 4922])
    headers = ("級別", "檢查項目", "狀態", "驗收結果／下一步")
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_fill(cell, ORANGE)
        add_text(cell, header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    for r_idx, row in enumerate(data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if c_idx in (0, 2):
                set_cell_fill(cell, CREAM if value not in ("待人工實測", "待外部條件") else ALERT)
                add_text(cell, value, bold=c_idx == 2, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_text(cell, value)


def add_page_inventory(doc: Document) -> None:
    data = [
        ("儀表板", "日常", "今日三件事優先；移除上線技術面板", "已改善"),
        ("撰寫公文", "日常", "保留兩步流程、儲存狀態與預覽", "已改善"),
        ("電子用印", "日常", "四步流程；整併合約用印", "已改善"),
        ("簽核進度", "日常", "用狀態籤快速篩選，手機單欄", "已改善"),
        ("公文收錄", "總務", "收錄／新增／派發分頁，手機單欄", "已改善"),
        ("公文查詢", "日常", "只保留關鍵字、類型、狀態", "已改善"),
        ("印章檔案庫", "授權角色", "權限提醒、單一 mm 尺寸、版本不可變", "已改善"),
        ("系統設定", "執行長", "不進一般主選單，手機正常捲動", "已改善"),
        ("通知中心", "次要", "寬內容限制在容器內", "已改善"),
        ("合約管理", "次要", "保留查詢；實際標記轉電子用印", "已改善"),
        ("合約用印舊網址", "相容", "不顯示於主選單，仍導向電子用印", "已改善"),
        ("流程控管", "次要", "Finance 固定流程說明；手機寬表格局部捲動", "已改善"),
        ("稽催追蹤", "次要", "手機單欄，不再撐開頁面", "已改善"),
        ("報表統計", "次要", "圖表／表格自身捲動，不帶動整頁", "已改善"),
        ("交換介接", "維運", "維持 Mock／停用，不讓一般使用者誤觸正式交換", "已確認"),
        ("帳號權限", "管理", "Finance 為唯一主檔；eDoc 顯示唯讀投影", "已確認"),
        ("維運中心", "維運", "不在日常主選單；長代碼可換行", "已改善"),
    ]
    table = doc.add_table(rows=1 + len(data), cols=4)
    set_table_geometry(table, [1900, 1300, 5572, 2000])
    for idx, header in enumerate(("頁面", "使用頻率", "化繁為簡方向", "目前狀態")):
        cell = table.cell(0, idx)
        set_cell_fill(cell, ORANGE)
        add_text(cell, header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    for r_idx, row in enumerate(data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx % 2 == 0:
                set_cell_fill(cell, CREAM)
            add_text(cell, value, bold=c_idx == 0, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx in (1, 3) else WD_ALIGN_PARAGRAPH.LEFT)


def add_revision_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=5)
    set_table_geometry(table, [1200, 1500, 4200, 1800, 2072])
    for idx, header in enumerate(("版次", "日期", "異動摘要", "起草", "審查／核准")):
        cell = table.cell(0, idx)
        set_cell_fill(cell, ORANGE)
        add_text(cell, header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    values = ("v0.1", "2026-08-24", "首次建立；納入桌機／手機逐頁檢視與 Before／After", "系統產製", "［待公司核定］")
    for idx, value in enumerate(values):
        add_text(table.cell(1, idx), value, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 2 else WD_ALIGN_PARAGRAPH.LEFT)


def set_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        for run in list(p.runs):
            p._p.remove(run._r)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run("歲悅長照｜總務 eDoc UIUX改善方案　　第 ")
        set_run_font(run, size=12, color=MUTED)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)
        end = p.add_run(" 頁")
        set_run_font(end, size=12, color=MUTED)


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(12)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.header_distance = Mm(2.2)
        section.footer_distance = Mm(3)
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Kaiti TC"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Kaiti TC")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Kaiti TC")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Kaiti TC")
        style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Kaiti TC")
    doc.styles["Normal"].font.size = Pt(12)
    doc.core_properties.title = "eDoc 電子公文用印系統 UI／UX 改善方案"
    doc.core_properties.subject = "桌機版與手機版逐頁檢視及 Before／After 改善說明"
    doc.core_properties.author = "歲悅長照（系統產製）"
    doc.core_properties.keywords = "eDoc, UIUX, 電子用印, 公文, 手機版, Before After"


def build() -> None:
    required = []
    for slug in ("dashboard", "compose", "electronicSeal", "approvalLog", "inbound", "search", "workflow", "seals"):
        for state in ("before", "after"):
            for device in ("desktop", "mobile"):
                required.append(SHOTS / f"{slug}-{state}-{device}.png")
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing screenshots: " + ", ".join(missing))

    doc = Document(str(TEMPLATE))
    clear_template_body(doc)
    configure_document(doc)
    set_footer(doc)

    # Editorial-cover structure with approved Suiyue visual tokens.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("系統上線前改善提案"), size=14, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("eDoc 電子公文用印系統"), size=28, bold=True, color=BROWN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_run_font(p.add_run("UI／UX 改善方案"), size=24, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    set_run_font(p.add_run("桌機版與手機版逐頁檢視｜Before / After｜v0.1"), size=13, color=MUTED)
    add_metadata_table(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    set_run_font(p.add_run("本文件使用去識別化的本機測試資料製作畫面，不含正式公文內容或印章原圖。"), size=12, color=MUTED)

    add_page_break(doc)
    add_heading(doc, "壹、執行摘要", 1)
    add_callout(doc, "核心結論", "系統最需要的不是增加更多功能，而是讓每個角色一進來只看到現在要做的事。第一輪已完成主選單縮減、首頁任務化、手機捲動修復、17 頁橫向跑版修復、44px 觸控區及中文化提示。", fill=CREAM)
    add_body(doc, "本次以一般使用者、簽核者、總務／行政與執行長四種視角，逐一檢視 17 個可進入頁面；其中 8 個核心頁面另附桌機與手機的 Before／After 實際畫面。")
    add_bullet(doc, "一般使用者：主選單集中在儀表板、撰寫公文、電子用印、簽核進度、通知與查詢。")
    add_bullet(doc, "簽核者：首頁直接顯示待簽案件；流程設定不再佔用日常入口。")
    add_bullet(doc, "總務／行政：保留公文收錄、印章檔案庫與全公司案件視角。")
    add_bullet(doc, "執行長：保留設定與印章權限，但維運、帳號與交換功能維持次要入口。")
    add_heading(doc, "一、已落地的 P0 改善與上線驗收狀態", 2)
    add_status_table(doc)

    add_page_break(doc)
    add_heading(doc, "貳、設計原則與化繁為簡方法", 1)
    principles = [
        ("1. 一頁一個主要任務", "每一頁只保留一個最明顯的下一步；其他動作降為次要按鈕或放在詳細區。"),
        ("2. 依角色縮減入口", "日常入口維持 6–8 個；設定、報表、稽催與維運不與日常工作並列。"),
        ("3. 先說人話再說系統話", "使用「今日待辦、兩步完成、四步完成、固定簽核規則」取代不必要的英文與技術名稱。"),
        ("4. 手機不能靠縮小桌機版", "內容改為單欄；寬表格在自己的區塊捲動，整頁永遠不橫移。"),
        ("5. 觸控優先", "按鈕與主要操作區至少 44px，避免誤觸與需要精準點擊。"),
        ("6. 字體用途清楚", "介面、導覽與表單一律易讀黑體；只有正式公文預覽、列印與 PDF 文字層使用標楷體。"),
        ("7. 保留相容、不保留重複", "合約用印舊網址可繼續使用，但主選單只顯示電子用印一個入口。"),
        ("8. 技術頁面不打擾日常", "Go／No-Go、帳號、交換、背景任務與維運資訊留在授權頁面，不放在使用者首頁。"),
    ]
    for title, body in principles:
        add_heading(doc, title, 3)
        add_body(doc, body)

    add_page_break(doc)
    add_heading(doc, "參、17 個頁面逐頁檢視", 1)
    add_body(doc, "下表列出執行長可進入的完整頁面集合。所有頁面均以 390×844 手機視窗實測：整頁橫向溢出為 0、內容可垂直捲動；主要操作區皆達 44px，僅原生勾選框維持平台標準尺寸。")
    add_page_inventory(doc)

    comparisons = [
        ("一、儀表板", "dashboard", "首頁從上線技術檢核改為今日三件事；技術檢核回到維運中心。"),
        ("二、撰寫公文", "compose", "保留兩步式填寫與確認；中文提示與可見儲存狀態降低不確定感。"),
        ("三、電子用印", "electronicSeal", "四步顯示進度；同一頁處理公文、合約與其他 PDF，避免雙入口。"),
        ("四、簽核進度", "approvalLog", "狀態籤、搜尋與案件清單先後清楚；手機維持單欄與可到達操作。"),
        ("五、公文收錄", "inbound", "收錄清單、新增收錄、內部派發分區；手機先看篩選與清單。"),
        ("六、公文查詢", "search", "只保留關鍵字、類型、狀態與查詢／重設，空結果清楚顯示。"),
        ("七、流程控管", "workflow", "不再是日常主選單；固定流程來源說明更明確，手機寬表格只在區塊內捲動。"),
        ("八、印章檔案庫", "seals", "權限與版本規則置頂；尺寸使用單一 mm 欄位，未授權者只能查看。"),
    ]
    add_page_break(doc)
    add_heading(doc, "肆、核心頁面 Before／After", 1)
    add_body(doc, "以下畫面均由本機隔離測試環境與去識別化資料產生。Before 為本次介面調整前，After 為本次調整後。")
    for idx, (heading, slug, note) in enumerate(comparisons):
        if idx:
            add_page_break(doc)
        add_heading(doc, heading, 2)
        add_body(doc, note)
        add_comparison(doc, "桌機版", SHOTS / f"{slug}-before-desktop.png", SHOTS / f"{slug}-after-desktop.png")
        add_comparison(doc, "手機版", SHOTS / f"{slug}-before-mobile.png", SHOTS / f"{slug}-after-mobile.png", mobile=True)

    add_page_break(doc)
    add_heading(doc, "伍、後續優化路線圖", 1)
    add_heading(doc, "一、P0｜上線前必驗", 2)
    add_bullet(doc, "由至少一位一般員工、一位課長／主任、一位總務與一位執行長，從模組頁實際點入 eDoc，逐筆記錄 5 秒內進入結果。")
    add_bullet(doc, "以正式權限走一次：建立公文→送簽→逐關下載附件→總務用印→申請人下載。")
    add_bullet(doc, "以 A4 直式與橫式 PDF 各走一次電子用印；印章素材未上傳前，只驗證流程與預覽，不宣稱正式用印完成。")
    add_bullet(doc, "Ready 檢查需為綠燈；防毒或外部維運條件若未完成，應維持 No-Go。")
    add_heading(doc, "二、P1｜上線後一週", 2)
    add_bullet(doc, "簽核進度增加「目前在誰手上、已等多久、下一步」三個人話欄位。")
    add_bullet(doc, "通知中心預設只顯示與本人有關的未讀與逾期；通知規則移到設定頁。")
    add_bullet(doc, "印章檔案庫將上傳、版本與使用紀錄拆成三個分頁，避免同頁操作過密。")
    add_heading(doc, "三、P2｜穩定後", 2)
    add_bullet(doc, "加入角色化新手導覽與空狀態範例，但不做長篇教學彈窗。")
    add_bullet(doc, "以實際使用數據檢查任務完成時間、退回率、找不到功能比例與手機使用比例。")
    add_bullet(doc, "正式電子公文交換仍待主管機關提供 jAgent／API／SDK／封包規格後另案實作；本階段維持 Mock／停用。")

    add_heading(doc, "六、驗收標準", 2)
    add_callout(doc, "可上線定義", "四角色真實登入成功、核心流程完整走通、附件可被後續簽核人查看與下載、手機 17 頁無整頁橫向跑版、Ready 檢查通過；公司印章正式素材可依已核定排除項目另行完成。", fill=BLUE)

    # Keep the final section on a fresh page without inserting an empty
    # page-break paragraph that can collide with the repeating header in Word.
    add_heading(doc, "陸、文件治理與修訂紀錄", 1, page_break_before=True)
    add_revision_table(doc)
    add_heading(doc, "一、來源與製作紀錄", 2)
    add_body(doc, f"公司 Word 範本：{TEMPLATE.name}；SHA-256：{sha256(TEMPLATE)}。")
    add_body(doc, "畫面來源：本機隔離測試環境，桌機 1440×1000、手機 390×844；截圖內容為測試資料，不含 Seal Vault 原圖。")
    add_body(doc, "處理模式：regenerate-from-approved-data；來源範本唯讀，輸出另存新檔。")
    add_heading(doc, "二、待公司核定欄位", 2)
    add_bullet(doc, "文件編號、正式管理單位、機密級別、保存年限、正式位置、審查者與核准者。")
    add_bullet(doc, "正式上線日期與本文件由審查中轉為已核准的版次。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
